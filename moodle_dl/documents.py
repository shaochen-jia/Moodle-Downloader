"""Writing a note in the format the reader actually opens.

Markdown is convenient for a developer and unfamiliar to everyone else, so
Word is the default and plain text the fallback. The same content is
rendered from one simple structure rather than written three times.
"""
from __future__ import annotations

import re
from pathlib import Path

FORMATS = ("docx", "txt", "md")
DEFAULT_FORMATS = ("docx", "txt")


def _clean(text: str) -> str:
    """Strip Markdown emphasis so plain readers do not see stray symbols."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def write(path_stem: Path, markdown: str, formats=DEFAULT_FORMATS) -> list[Path]:
    """Write the same note in each requested format. Returns the files made.

    Formats no longer wanted are removed, so switching to Word does not
    leave the old Markdown copy sitting beside it looking current.
    """
    for fmt in FORMATS:
        if fmt not in formats:
            stale = path_stem.with_suffix(f".{fmt}")
            if stale.exists():
                try:
                    stale.unlink()
                except OSError:
                    pass

    written: list[Path] = []
    for fmt in formats:
        if fmt == "md":
            p = path_stem.with_suffix(".md")
            p.write_text(markdown, encoding="utf-8")
            written.append(p)
        elif fmt == "txt":
            p = path_stem.with_suffix(".txt")
            p.write_text(_to_text(markdown), encoding="utf-8")
            written.append(p)
        elif fmt == "docx":
            p = _to_docx(path_stem.with_suffix(".docx"), markdown)
            if p:
                written.append(p)
    return written


def _to_text(markdown: str) -> str:
    out = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = _clean(stripped.lstrip("#").strip())
            out += ["", title, "-" * len(title)]
        elif stripped.startswith(("- ", "* ")):
            out.append("  * " + _clean(stripped[2:]))
        elif stripped.startswith("> "):
            out.append("  | " + _clean(stripped[2:]))
        elif stripped == "---":
            out.append("")
        else:
            out.append(_clean(line))
    return "\n".join(out).strip() + "\n"


def _to_docx(path: Path, markdown: str) -> Path | None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return None
    try:
        doc = Document()
        in_code = False
        for raw in markdown.splitlines():
            line = raw.rstrip()
            stripped = line.strip()
            if stripped.startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                p = doc.add_paragraph(line)
                p.runs[0].font.name = "Consolas"
                p.runs[0].font.size = Pt(9)
                continue
            if not stripped or stripped == "---":
                continue
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                doc.add_heading(_clean(stripped.lstrip("#").strip()),
                                level=min(level, 4))
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(_clean(stripped[2:]), style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(_clean(re.sub(r"^\d+\.\s", "", stripped)),
                                  style="List Number")
            elif stripped.startswith("> "):
                p = doc.add_paragraph(_clean(stripped[2:]))
                p.runs[0].italic = True
            elif stripped.startswith("*") and stripped.endswith("*"):
                p = doc.add_paragraph(_clean(stripped))
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(9)
            else:
                doc.add_paragraph(_clean(stripped))
        doc.save(str(path))
        return path
    except Exception as e:
        print(f"    ! Word version of {path.name}: {type(e).__name__}: {e}")
        return None
